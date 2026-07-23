#!/usr/bin/env python3
"""Build a Google Docs-targeted DOCX manuscript from the climate revision draft."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_MD = ROOT / "docs/manuscripts/fire_vase_developmental_morphology/manuscript_climate_revision_science_style.md"
LEGENDS_MD = ROOT / "figures/climate_revision_main/figure_legends.md"
FIGURE_DIR = ROOT / "figures/climate_revision_main"
OUTPUT = ROOT / "output/docx/fire_vase_climate_revision_google_docs.docx"


def set_cell_margins(table, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, RGBColor(0, 0, 0)),
        ("Heading 2", 16, 18, 6, RGBColor(0, 0, 0)),
        ("Heading 3", 14, 16, 4, RGBColor(67, 67, 67)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for name in ["Caption", "Subtitle"]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.line_spacing = 1.15

    doc.styles["Caption"].font.size = Pt(10)
    doc.styles["Caption"].font.bold = False
    doc.styles["Caption"].font.italic = False
    doc.styles["Caption"].font.color.rgb = RGBColor(85, 85, 85)
    doc.styles["Subtitle"].font.size = Pt(11)
    doc.styles["Subtitle"].font.color.rgb = RGBColor(85, 85, 85)

    ref = doc.styles.add_style("Reference", 1)
    ref.base_style = normal
    ref.font.name = "Arial"
    ref._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    ref.font.size = Pt(10)
    ref.paragraph_format.left_indent = Inches(0.25)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.space_after = Pt(8)
    ref.paragraph_format.line_spacing = 1.15


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)


def parse_legends() -> dict[int, tuple[str, str]]:
    legends: dict[int, tuple[str, str]] = {}
    current_num: int | None = None
    current_title = ""
    buf: list[str] = []
    for raw in LEGENDS_MD.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        match = re.match(r"## Fig\. (\d+)\. (.+)", line)
        if match:
            if current_num is not None:
                legends[current_num] = (current_title, " ".join(buf).strip())
            current_num = int(match.group(1))
            current_title = match.group(2)
            buf = []
        elif line and current_num is not None:
            buf.append(line)
    if current_num is not None:
        legends[current_num] = (current_title, " ".join(buf).strip())
    return legends


def add_paragraph_with_basic_emphasis(doc: Document, text: str, style: str = "Normal") -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(text)


def add_markdown_body(doc: Document) -> None:
    lines = MANUSCRIPT_MD.read_text(encoding="utf-8").splitlines()
    in_references = False
    title_done = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            if not title_done:
                add_title(doc, line[2:])
                title_done = True
            else:
                doc.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            heading = line[3:]
            in_references = heading == "References and Notes"
            doc.add_heading(heading, level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("Authors:") or line.startswith("Correspondence:"):
            add_paragraph_with_basic_emphasis(doc, line, "Subtitle")
            continue
        if in_references and re.match(r"^\d+\.\s", line):
            add_paragraph_with_basic_emphasis(doc, line, "Reference")
            continue
        add_paragraph_with_basic_emphasis(doc, line)


def add_figures(doc: Document) -> None:
    legends = parse_legends()
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for num in range(1, 6):
        title, caption = legends[num]
        if num > 1:
            doc.add_page_break()
        image = FIGURE_DIR / f"Figure_{num}_climate_revision.png"
        doc.add_picture(str(image), width=Inches(6.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p = doc.add_paragraph(style="Caption")
        label = caption_p.add_run(f"Fig. {num}. {title} ")
        label.bold = True
        caption_p.add_run(caption)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_markdown_body(doc)
    add_figures(doc)
    for table in doc.tables:
        set_cell_margins(table)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
